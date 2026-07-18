"""
Word Document Parser — Handles .doc and .docx files.
Extracts tables and text-based tabular data.
"""
from __future__ import annotations

import logging
import re
import time
from pathlib import Path
from typing import Optional

from backend.models import Cell, Table, ProcessingResult, ProcessingStatus
from backend.utils import detect_data_type

logger = logging.getLogger(__name__)


class WordParser:
    """Parses Word documents (.docx, .doc) extracting tables and structured text."""

    def parse(self, file_path: str | Path) -> ProcessingResult:
        """Parse a Word document."""
        result = ProcessingResult()
        result.filename = Path(file_path).name
        result.file_category = "word"
        result.page_count = 1
        start_time = time.time()

        ext = Path(file_path).suffix.lower()

        try:
            result.status = ProcessingStatus.OCR_IN_PROGRESS

            if ext == ".docx":
                tables = self._parse_docx(str(file_path))
            elif ext == ".doc":
                tables = self._parse_doc(str(file_path))
            else:
                result.mark_failed(f"Unsupported Word format: {ext}")
                return result

            result.tables = tables
            result.ocr_engine_used = "python-docx" if ext == ".docx" else "textract"
            result.status = ProcessingStatus.READY_FOR_REVIEW
            result.processing_time = time.time() - start_time

            if not tables:
                result.warnings.append("No tables found in the document")

        except Exception as e:
            result.mark_failed(f"Word parsing failed: {str(e)}")
            logger.error(f"Word parsing error: {e}", exc_info=True)

        return result

    def _parse_docx(self, file_path: str) -> list[Table]:
        """Parse .docx using python-docx."""
        from docx import Document as DocxDocument

        tables: list[Table] = []
        doc = DocxDocument(file_path)

        # Extract tables
        for tbl_idx, docx_table in enumerate(doc.tables):
            table = Table()
            table.source_engine = "python-docx"
            table.title = f"Table {tbl_idx + 1}"

            for row_idx, row in enumerate(docx_table.rows):
                cells: list[Cell] = []
                for col_idx, docx_cell in enumerate(row.cells):
                    cell = Cell(row=row_idx, col=col_idx)
                    text = docx_cell.text.strip()
                    cell.raw_value = text
                    cell.value = text
                    cell.confidence = 1.0

                    # Check for merged cells
                    if hasattr(docx_cell, '_tc'):
                        tc = docx_cell._tc
                        grid_span = tc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                        if grid_span is not None:
                            cell.col_span = int(grid_span.get(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1'))
                            cell.is_merged = cell.col_span > 1

                    # Check formatting
                    for paragraph in docx_cell.paragraphs:
                        for run in paragraph.runs:
                            if run.bold:
                                cell.font_bold = True
                            if run.font.size:
                                cell.font_size = run.font.size.pt

                    # Detect data type
                    if text:
                        dt, parsed, fmt = detect_data_type(text)
                        cell.data_type = dt
                        if dt in ("number", "currency", "percentage"):
                            cell.value = parsed
                        cell.format_string = fmt

                    if row_idx == 0:
                        cell.is_header = True
                        cell.font_bold = True

                    cells.append(cell)

                table.cells.append(cells)

            if table.cells:
                table.headers = [str(c.value) for c in table.cells[0]]
                table.num_rows = len(table.cells)
                table.num_cols = max(len(r) for r in table.cells) if table.cells else 0
                tables.append(table)

        # Also extract text-based tables from paragraphs
        text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text_content and not tables:
            text_table = self._extract_text_table(text_content)
            if text_table:
                tables.append(text_table)

        return tables

    def _parse_doc(self, file_path: str) -> list[Table]:
        """Parse legacy .doc files. Try multiple methods."""
        # Method 1: Try antiword/textract
        try:
            import subprocess
            text = subprocess.check_output(
                ["antiword", file_path],
                timeout=30,
                stderr=subprocess.DEVNULL
            ).decode("utf-8", errors="ignore")

            if text.strip():
                table = self._extract_text_table(text)
                return [table] if table else []
        except (FileNotFoundError, subprocess.SubprocessError):
            pass

        # Method 2: Try python-docx (sometimes works with .doc)
        try:
            return self._parse_docx(file_path)
        except Exception:
            pass

        logger.warning(f"Could not parse .doc file: {file_path}")
        return []

    def _extract_text_table(self, text: str) -> Optional[Table]:
        """Extract table from plain text with tabular structure."""
        lines = [line for line in text.split('\n') if line.strip()]
        if len(lines) < 2:
            return None

        table = Table()
        table.source_engine = "text_extraction"
        table.has_borders = False

        for line_idx, line in enumerate(lines):
            parts = re.split(r'\s{2,}|\t|\|', line)
            parts = [p.strip() for p in parts if p.strip()]

            if not parts:
                continue

            cells: list[Cell] = []
            for col_idx, part in enumerate(parts):
                cell = Cell(row=line_idx, col=col_idx)
                cell.raw_value = part
                cell.value = part
                cell.confidence = 1.0

                if part:
                    dt, parsed, fmt = detect_data_type(part)
                    cell.data_type = dt
                    if dt in ("number", "currency", "percentage"):
                        cell.value = parsed
                    cell.format_string = fmt

                if line_idx == 0:
                    cell.is_header = True
                    cell.font_bold = True

                cells.append(cell)

            table.cells.append(cells)

        if table.cells and len(table.cells) > 1:
            table.headers = [str(c.value) for c in table.cells[0]]
            table.num_rows = len(table.cells)
            table.num_cols = max(len(r) for r in table.cells) if table.cells else 0
            return table

        return None
